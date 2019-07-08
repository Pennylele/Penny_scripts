#! /usr/bin/env python3
# SRT2DOCX.py is used for writing content from a SRT file into a DOCX file.

from docx import Document
import io, os, re
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Inches


source = input("What's the source SRT? ")
target = input("What's the target SRT? ")

f = io.open(source, encoding="utf-8", mode="r+")
regexContent = re.compile('(\d{2}:\d{2}:\d{2},\d{3})\s-->')
StartTime = regexContent.findall(f.read())
document = Document()


f = io.open(source, encoding="utf-8", mode="r+")
sourceCodes = re.findall(r"\d+\n\d\d:\d\d:\d+,\d+\s-->\s\d\d:\d\d:\d+,\d+\n", f.read())
subNo = len(sourceCodes)
f = io.open(source, encoding="utf-8", mode="r+")
subtitles = re.findall(r":\d+,\d+\n(.*?)\n\n", f.read(), re.DOTALL)
f = io.open(source, encoding="utf-8", mode="r+")
LastLine = re.findall(str(subNo) + r"\n\d\d:\d\d:\d+,\d+\s-->\s\d\d:\d\d:\d+,\d+\n(.*)", f.read(), re.DOTALL)
subtitles.append(LastLine[-1])
dictionary = dict(zip(StartTime, subtitles))

f2 = io.open(target, encoding="utf-8", mode="r+")
subtitles_target = re.findall(r":\d+,\d+\n(.*?)\n\n", f2.read(), re.DOTALL)
f2 = io.open(target, encoding="utf-8", mode="r+")
LastLine_target = re.findall(str(subNo) + r"\n\d\d:\d\d:\d+,\d+\s-->\s\d\d:\d\d:\d+,\d+\n(.*)", f2.read(), re.DOTALL)
subtitles_target.append(LastLine_target[-1])


table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
cell1 = hdr_cells[0]
cell1.text = 'Scene'
run = cell1.paragraphs[0].runs[0]
run.font.color.rgb = RGBColor(255, 255, 255)
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

cell2 = hdr_cells[1]
cell2.text = 'Start Time'
run = cell2.paragraphs[0].runs[0]
run.font.color.rgb = RGBColor(255, 255, 255)
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
cell2 = hdr_cells[2]
cell2.text = 'Audio'
run = cell2.paragraphs[0].runs[0]
run.font.color.rgb = RGBColor(255, 255, 255)
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
#hdr_cells[3].text = 'Translation'
cell3 = hdr_cells[3]
cell3.text = 'Audio Translation'
run = cell2.paragraphs[0].runs[0]
run.font.color.rgb = RGBColor(255, 255, 255)
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
table.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

r = 1
for StartTime, subtitle in sorted(dictionary.items()):
    row_cells = table.add_row().cells
    row_cells[0].text = str(r)
    row_cells[1].text = str(StartTime)
    row_cells[2].text = str(subtitle)
    row_cells[3].text = str(subtitles_target[r-1])
    r += 1

hdr_cells[0].width = Cm(0)
hdr_cells[2].width = Inches(5)
hdr_cells[3].width = Inches(10)

document.save(target + '.docx')

for file in os.listdir('.'):
    filename = os.path.join(os.getcwd(), file)
    newname = filename.replace('.srt.docx', '.docx')
    os.rename(filename, newname)

